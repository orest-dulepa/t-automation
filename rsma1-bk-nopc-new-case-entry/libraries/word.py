import subprocess
from os.path import dirname
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document
import libraries.common as com
import re
import os


def export_to_pdf(path: str, outpath: str):
    exec_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
    new_file_abs = os.path.abspath(path)
    new_file_abs = re.sub(r'\.\w+$', '.pdf', new_file_abs)
    out_dir = os.path.abspath(dirname(path))
    try:
        a = subprocess.run(['cmd', '/c', exec_path, '--headless', '--convert-to', 'pdf', '--outdir', out_dir, path],
                           capture_output=True, text=True)
        com.log_message(a.stdout, 'TRACE')
        com.log_message(a.stderr, 'TRACE')
    except Exception as ex:
        com.log_message('Failed to convert file', 'TRACE')
        com.log_message(ex, 'TRACE')
        com.log_message(ex.with_traceback(None), 'TRACE')
    os.rename(new_file_abs, outpath)


def save_as_docx(path: str) -> str:
    exec_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
    new_file_abs = os.path.abspath(path)
    new_file_abs = re.sub(r'\.\w+$', '.docx', new_file_abs)
    out_dir = os.path.abspath(dirname(path))
    try:
        a = subprocess.run(['cmd', '/c', exec_path, '--headless', '--convert-to', 'docx', '--outdir', out_dir, path],
                           capture_output=True, text=True)
        com.log_message(a.stdout, 'TRACE')
        com.log_message(a.stderr, 'TRACE')
    except Exception as ex:
        com.log_message('Failed to convert file', 'TRACE')
        com.log_message(ex, 'TRACE')
        com.log_message(ex.with_traceback(None), 'TRACE')
    return new_file_abs


def checked_element():
    elm = OxmlElement('w:checked')
    elm.set(qn('w:val'), "true")
    return elm


def yes_no_check(doc, yes_no, tableIdx, coords):
    com.log_message(coords, yes_no, 'TRACE')
    if yes_no == 'n':
        index = 0
        x = doc.tables[tableIdx].cell(coords[0], coords[1])._element.xpath('.//w:checkBox')
        x[index].append(checked_element())
    elif yes_no == 'y':
        index = 1
        x = doc.tables[tableIdx].cell(coords[0], coords[1])._element.xpath('.//w:checkBox')
        x[index].append(checked_element())
    else:
        com.log_message("value was neither yes or no", 'TRACE')
        pass


def process_document(path: str, cb_state1: str, cb_state2: str, cb_state3: str):
    new_path = save_as_docx(path)
    doc = Document(new_path)
    yes_no_check(doc, cb_state1, 1, [4, 1])
    yes_no_check(doc, cb_state2, 2, [1, 1])
    yes_no_check(doc, cb_state3, 2, [2, 1])
    a = doc.tables[2].cell(2, 1).paragraphs[4].text
    doc.tables[2].cell(2, 1).paragraphs[4].text = re.sub(r'(New mortgage payment:\t\$\s.+)',
                                                         'New mortgage payment:	$ 		', a)
    output_path_w = os.path.join(os.path.dirname(new_path), 'nopc.docx')
    output_path_pdf = os.path.join(os.path.dirname(new_path), 'Checked_NOPC.pdf')
    doc.save(output_path_w)

    export_to_pdf(output_path_w, output_path_pdf)
    os.remove(path)
    os.remove(new_path)
    os.remove(output_path_w)
    return output_path_pdf
